"""
Genera un file DOCX da ExtractionResult.
Usa python-docx — nessuna dipendenza Java, nessun JAR.
Gestisce: elementi strutturati, markdown con tabelle, separatori tra certificati.
"""
from __future__ import annotations
from pathlib import Path
import logging
import re

from src.extractor_pymupdf import ExtractionResult

logger = logging.getLogger(__name__)


def to_docx(result: ExtractionResult, output_path: Path) -> Path:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as exc:
        raise ImportError("python-docx non installato. Esegui: pip install python-docx") from exc

    doc = Document()

    # ── Stile base ────────────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # Margini stretti per tabelle più larghe
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    # ── Intestazione ─────────────────────────────────────────────────────────
    info = doc.add_paragraph()
    info.add_run(f"Motore: {result.engine}  |  Pagine: {result.page_count}").italic = True
    info.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if result.error:
        doc.add_paragraph(f"⚠ Errore: {result.error}")

    # ── Scelta renderer ───────────────────────────────────────────────────────
    # Se il markdown contiene tabelle Markdown → renderer specializzato
    if result.markdown and "|---|" in result.markdown:
        _render_cert_markdown(doc, result.markdown)
    elif result.elements:
        _render_elements(doc, result)
    elif result.markdown:
        _render_markdown_simple(doc, result.markdown)
    else:
        doc.add_paragraph("Nessun contenuto estratto dal PDF.")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("DOCX creato: %s (%d bytes)", output_path, output_path.stat().st_size)
    return output_path


# ── Renderer per certificati SSN (tabelle Markdown) ───────────────────────────
def _render_cert_markdown(doc, markdown: str) -> None:
    """Converte markdown con tabelle in DOCX con tabelle Word."""
    from docx.shared import RGBColor, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    BLU      = RGBColor(0x1F, 0x49, 0x7D)
    BLU_LT   = "D6E4F0"
    GRIGIO   = RGBColor(0x59, 0x59, 0x59)

    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()

        # Titolo ## (certificato N)
        if line.startswith("## "):
            doc.add_paragraph()
            h = doc.add_heading(line[3:], level=2)
            for run in h.runs:
                run.font.color.rgb = BLU
            i += 1
            continue

        # Titolo # 
        if line.startswith("# "):
            h = doc.add_heading(line[2:], level=1)
            for run in h.runs:
                run.font.color.rgb = BLU
            i += 1
            continue

        # Separatore ---
        if line == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(6)
            pPr = p._p.get_or_add_pPr()
            from docx.oxml import OxmlElement
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"),   "single")
            bottom.set(qn("w:sz"),    "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "1F497D")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Inizio tabella Markdown: riga con |
        if line.startswith("|") and i + 1 < len(lines) and "|---|" in lines[i+1]:
            # Raccoglie tutte le righe della tabella
            table_lines = []
            j = i
            while j < len(lines) and lines[j].strip().startswith("|"):
                table_lines.append(lines[j].strip())
                j += 1
            _add_md_table(doc, table_lines, BLU_LT, BLU)
            i = j
            continue

        # Riga vuota
        if not line:
            i += 1
            continue

        # Testo normale
        doc.add_paragraph(line)
        i += 1


def _add_md_table(doc, table_lines: list[str], header_fill: str, header_color) -> None:
    """Converte righe Markdown table in tabella Word."""
    from docx.shared import RGBColor, Pt
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Filtra riga separatore |---|
    rows = [l for l in table_lines if not re.match(r"^\|[-| :]+\|$", l)]
    if not rows:
        return

    # Parse celle
    def parse_row(line):
        cells = [c.strip() for c in line.strip("|").split("|")]
        return cells

    parsed = [parse_row(r) for r in rows]
    n_cols = max(len(r) for r in parsed)

    # Larghezza colonne in twips (1 cm = 567 twips)
    # Pagina A4 con margini 2cm: ~16.7cm disponibili = 9450 twips circa
    col_widths = _calc_col_widths(n_cols)

    table = doc.add_table(rows=0, cols=n_cols)
    table.style = "Table Grid"

    for r_idx, cells in enumerate(parsed):
        row = table.add_row()
        is_header = (r_idx == 0)
        for c_idx in range(n_cols):
            cell = row.cells[c_idx]
            text = cells[c_idx] if c_idx < len(cells) else ""
            # Rimuove bold markdown **testo**
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)

            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(10)

            if is_header:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                _set_cell_bg(cell, "1F497D")
            else:
                # Colonna label (prima col) — sfondo azzurro chiaro
                if c_idx == 0:
                    run.bold = True
                    _set_cell_bg(cell, header_fill)

            # Padding cella
            _set_cell_margins(cell)
            # Larghezza cella
            _set_cell_width(cell, col_widths[c_idx])


def _calc_col_widths(n_cols: int) -> list[int]:
    """Calcola larghezze colonne in twips per pagina A4 con margini 2cm."""
    total = 9360  # twips disponibili (~16.5cm)
    if n_cols == 2:
        return [2800, 6560]
    if n_cols == 3:
        return [2000, 4000, 3360]
    each = total // n_cols
    return [each] * n_cols


def _set_cell_bg(cell, hex_color: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top",top),("bottom",bottom),("left",left),("right",right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"),    str(val))
        el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def _set_cell_width(cell, width_twips: int) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(width_twips))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


# ── Renderer elementi strutturati ─────────────────────────────────────────────
def _render_elements(doc, result: ExtractionResult) -> None:
    from docx.shared import RGBColor, Pt
    BLU = RGBColor(0x44, 0x72, 0xC4)
    current_page = -1
    for elem in result.elements:
        if elem.page != current_page:
            current_page = elem.page
            sep = doc.add_paragraph(f"── Pagina {elem.page + 1} ──")
            sep.runs[0].bold = True
            sep.runs[0].font.color.rgb = BLU
        if elem.type == "heading":
            doc.add_heading(elem.text, level=max(1, min(elem.level, 9)))
        elif elem.type == "image":
            p = doc.add_paragraph()
            run = p.add_run("[immagine]")
            run.italic = True
        else:
            doc.add_paragraph(elem.text)


# ── Renderer markdown semplice ────────────────────────────────────────────────
def _render_markdown_simple(doc, markdown: str) -> None:
    for line in markdown.splitlines():
        line = line.strip()
        if not line or line.startswith("---"):
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        else:
            doc.add_paragraph(line)
